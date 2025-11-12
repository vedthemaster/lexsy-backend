from typing import BinaryIO
import os
from docx import Document as DocxDocument

from api.v2.models.models import Document
from api.v2.repository.document_repository import document_repo_ins


class DocumentGeneratorService:
    """Service for generating filled documents"""

    async def generate_document(self, document_id: str) -> str:
        """Generate a filled document and return the file path"""

        document = await document_repo_ins.find_by_id(document_id)
        if not document:
            raise ValueError(f"Document {document_id} not found")

        unfilled_count = sum(1 for p in document.placeholders if p.value is None)
        if unfilled_count > 0:
            raise ValueError(f"Document has {unfilled_count} unfilled placeholders")

        doc_path = document.temp_path if document.temp_path else document.path
        doc = DocxDocument(doc_path)

        for placeholder in document.placeholders:
            if placeholder.value is not None:
                for para in doc.paragraphs:
                    if placeholder.unique_marker in para.text:
                        for run in para.runs:
                            if placeholder.unique_marker in run.text:
                                run.text = run.text.replace(
                                    placeholder.unique_marker, str(placeholder.value)
                                )

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if placeholder.unique_marker in cell.text:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if placeholder.unique_marker in run.text:
                                            run.text = run.text.replace(
                                                placeholder.unique_marker,
                                                str(placeholder.value),
                                            )

        generated_dir = "uploads/generated"
        os.makedirs(generated_dir, exist_ok=True)

        generated_filename = f"filled_{document.title}"
        generated_path = os.path.join(generated_dir, generated_filename)

        doc.save(generated_path)

        return generated_path

    async def get_document_stream(self, document_id: str) -> tuple[BinaryIO, str]:
        """Get document as a stream for download"""

        generated_path = await self.generate_document(document_id)

        file_stream = open(generated_path, "rb")
        filename = os.path.basename(generated_path)

        return file_stream, filename


document_generator_service = DocumentGeneratorService()
