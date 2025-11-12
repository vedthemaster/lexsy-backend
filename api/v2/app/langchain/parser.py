from typing import List
import json
import uuid
import os
from docx import Document as DocxDocument
from langchain_openai import ChatOpenAI

from api.v2.app.langchain.tools import (
    PlaceholderDetectorTool,
    ContextAnalyzerTool,
    PlaceholderClassifierTool,
)
from api.v2.models.models import PlaceHolder, PlaceholderAnalysis, PlaceholderType
from config import config


class LangChainParser:
    """Document parser using LangChain tools for multi-stage analysis"""

    def __init__(self):
        self.llm = ChatOpenAI(
            model="gpt-4o-mini", temperature=0, api_key=config.OPENAI_API_KEY
        )
        self.detector_tool = PlaceholderDetectorTool()
        self.context_analyzer_tool = ContextAnalyzerTool()
        self.classifier_tool = PlaceholderClassifierTool()

    async def parse_document(self, file_path: str) -> tuple[List[PlaceHolder], str]:
        """Parse a document and extract placeholders with full analysis

        Returns:
            tuple: (placeholders, temp_document_path)
        """

        doc = DocxDocument(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        detections = self.detector_tool._run(full_text)
        print(f"Found {len(detections)} placeholders")

        placeholders = []
        for idx, detection in enumerate(detections):
            try:
                context_input = json.dumps(
                    {
                        "placeholder": detection.text,
                        "context_before": detection.context_before,
                        "context_after": detection.context_after,
                    }
                )
                context_analysis = self.context_analyzer_tool._run(context_input)

                classifier_input = json.dumps(
                    {
                        "placeholder": detection.text,
                        "semantic_meaning": context_analysis.semantic_meaning,
                        "context": f"{detection.context_before} ... {detection.context_after}",
                    }
                )
                classification = self.classifier_tool._run(classifier_input)

                cleaned_label = detection.text.strip("[]{}()<>_")

                if not cleaned_label or cleaned_label.strip() == "":
                    cleaned_label = self._extract_name_from_context(
                        detection.context_before,
                        detection.context_after,
                        context_analysis.semantic_meaning,
                    )

                name = cleaned_label

                question_hint = self._generate_question_hint(
                    detection.text,
                    name,
                    context_analysis.semantic_meaning,
                    classification.type,
                )

                validation_rules = (
                    [context_analysis.validation_hints]
                    if context_analysis.validation_hints
                    else []
                )

                analysis = PlaceholderAnalysis(
                    context_before=detection.context_before,
                    context_after=detection.context_after,
                    inferred_type=classification.type,
                    confidence_score=min(
                        detection.confidence, classification.confidence
                    ),
                    validation_rules=validation_rules,
                    suggested_value=None,
                    related_placeholders=[],
                    question_hint=question_hint,
                )

                unique_marker = f"{{{{PLACEHOLDER_{uuid.uuid4().hex[:8].upper()}}}}}"

                placeholder = PlaceHolder(
                    name=name,
                    placeholder=detection.text,
                    unique_marker=unique_marker,
                    regex=self._generate_regex_pattern(detection.text),
                    value=None,
                    analysis=analysis,
                )

                placeholders.append(placeholder)

            except Exception as e:
                print(f"Error analyzing placeholder {detection.text}: {e}")
                unique_marker = f"{{{{PLACEHOLDER_{uuid.uuid4().hex[:8].upper()}}}}}"
                cleaned_label = detection.text.strip("[]{}()<>_")
                placeholder = PlaceHolder(
                    name=cleaned_label,
                    placeholder=detection.text,
                    unique_marker=unique_marker,
                    regex=self._generate_regex_pattern(detection.text),
                    value=None,
                    analysis=None,
                )
                placeholders.append(placeholder)

        temp_doc_path = self._create_temp_document(file_path, placeholders)

        return placeholders, temp_doc_path

    def _generate_regex_pattern(self, placeholder_text: str) -> str:
        """Generate a regex pattern to match the placeholder in the document"""
        import re

        escaped = re.escape(placeholder_text)
        return escaped

    def _create_temp_document(
        self, original_path: str, placeholders: List[PlaceHolder]
    ) -> str:
        """Create a temporary document with unique markers replacing original placeholders

        Process each placeholder one at a time, replacing only the FIRST occurrence,
        then save and reload for the next placeholder. This handles duplicate placeholders correctly.
        """
        import re

        temp_dir = "uploads/temp"
        os.makedirs(temp_dir, exist_ok=True)

        original_filename = os.path.basename(original_path)
        temp_filename = f"temp_{uuid.uuid4().hex[:8]}_{original_filename}"
        temp_path = os.path.join(temp_dir, temp_filename)

        doc = DocxDocument(original_path)
        doc.save(temp_path)

        for idx, placeholder in enumerate(placeholders):
            doc = DocxDocument(temp_path)
            replaced = False

            pattern = re.compile(re.escape(placeholder.placeholder))

            if not replaced:
                for para in doc.paragraphs:
                    if pattern.search(para.text):
                        full_text = "".join(run.text for run in para.runs)

                        new_text = pattern.sub(
                            placeholder.unique_marker, full_text, count=1
                        )

                        if new_text != full_text:
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
                            else:
                                para.add_run(new_text)

                            replaced = True
                            break

            if not replaced:
                for table in doc.tables:
                    if replaced:
                        break
                    for row in table.rows:
                        if replaced:
                            break
                        for cell in row.cells:
                            if replaced:
                                break
                            for para in cell.paragraphs:
                                if pattern.search(para.text):
                                    full_text = "".join(run.text for run in para.runs)

                                    new_text = pattern.sub(
                                        placeholder.unique_marker, full_text, count=1
                                    )

                                    if new_text != full_text:
                                        for run in para.runs:
                                            run.text = ""
                                        if para.runs:
                                            para.runs[0].text = new_text
                                        else:
                                            para.add_run(new_text)

                                        replaced = True
                                        break

            if replaced:
                doc.save(temp_path)
                print(
                    f"Replaced placeholder {idx + 1}/{len(placeholders)}: {placeholder.placeholder} -> {placeholder.unique_marker}"
                )
            else:
                print(
                    f"Warning: Could not find placeholder {idx + 1}/{len(placeholders)}: {placeholder.placeholder}"
                )

        print(f"Created temp document: {temp_path}")
        return temp_path

    def _extract_name_from_context(
        self, context_before: str, context_after: str, semantic_meaning: str
    ) -> str:
        """Extract a meaningful name from context when placeholder is blank"""

        # Use LLM to extract concise name from context
        prompt = f"""Extract a short, descriptive name (2-4 words) for a blank placeholder based on context.

CONTEXT BEFORE: ...{context_before[-100:]}
[BLANK PLACEHOLDER]
CONTEXT AFTER: {context_after[:100]}...

SEMANTIC MEANING: {semantic_meaning}

Extract what this blank represents. Be specific and concise.

EXAMPLES:
Context: "payment by [___]" → "Investor Name"
Context: "of $[___]" → "Purchase Amount"
Context: "Date: [___]" → "Date"
Context: "Company Name: [___]" → "Company Name"

Return ONLY the name (2-4 words):"""

        try:
            name = self.llm.invoke(prompt).content.strip()
            # Clean up the name
            name = name.strip("\"'").strip()
            return name if name else "Required Information"
        except Exception:
            return "Required Information"

    def _generate_question_hint(
        self,
        placeholder: str,
        placeholder_name: str,
        semantic_meaning: str,
        placeholder_type: PlaceholderType,
    ) -> str:
        """Generate a helpful question hint for users"""

        prompt = f"""Create a clear question to ask a user to fill in this information.

Placeholder Text: {placeholder}
Field Name: {placeholder_name}
Meaning: {semantic_meaning}
Type: {placeholder_type.value}

Generate a natural, professional question asking for the "{placeholder_name}". 
Include format hints if needed (e.g., for dates, emails).

Return ONLY the question:"""

        try:
            return self.llm.invoke(prompt).content.strip()
        except Exception:
            return (
                f"What is the {placeholder_name}?"
                if placeholder_name
                else "Please provide the required information."
            )
