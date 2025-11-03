import os
import re
import uuid
from docx import Document as DocxDocument
from fastapi import HTTPException, status

from repository.document_repository import document_repo_ins


class DocumentGeneratorService:

    def __init__(self):
        self.output_dir = "uploads/generated"
        os.makedirs(self.output_dir, exist_ok=True)

    async def generate_filled_document(self, document_id: str):
        """
        Generate a new document with all placeholder values filled in

        Args:
            document_id: The ID of the document to generate

        Returns:
            Returns path to generated document file
        """
        document = await document_repo_ins.get_document_by_id(document_id)

        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document with id {document_id} not found",
            )

        unfilled = [ph.name for ph in document.placeholders if not ph.value]
        if unfilled:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Not all placeholders are filled. Missing: {', '.join(unfilled)}",
            )

        try:
            doc = DocxDocument(document.path)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error loading document: {str(e)}",
            )

        replacements_made = await self._replace_placeholders_in_doc(
            doc, document.placeholders
        )

        output_filename = f"filled_{uuid.uuid4().hex}_{document.title}"
        output_path = os.path.join(self.output_dir, output_filename)

        doc.save(output_path)

        return {
            "document_id": document_id,
            "title": document.title,
            "output_path": output_path,
            "output_filename": output_filename,
            "replacements_made": replacements_made,
        }

    async def _replace_placeholders_in_doc(
        self, doc: DocxDocument, placeholders: list
    ) -> int:
        """
        Replace placeholders in the document using simple regex patterns.
        Each placeholder is replaced only ONCE - the first occurrence found.
        Once a replacement is made, we move to the next placeholder.
        Returns the number of replacements made.
        """
        replacements_count = 0

        for placeholder in placeholders:
            if not placeholder.value:
                continue

            regex_pattern = placeholder.regex
            replacement_value = str(placeholder.value)
            replaced = False  # Track if we've replaced this placeholder

            # Try to replace in paragraphs first
            if not replaced:
                for paragraph in doc.paragraphs:
                    if replaced:
                        break

                    if paragraph.text:
                        original_text = paragraph.text
                        try:
                            # Check if pattern exists in this paragraph
                            if re.search(regex_pattern, original_text):
                                # Replace only the FIRST occurrence
                                new_text = re.sub(
                                    regex_pattern,
                                    replacement_value,
                                    original_text,
                                    count=1,
                                )

                                # Update paragraph and mark as replaced
                                self._update_paragraph_text(paragraph, new_text)
                                replacements_count += 1
                                replaced = True
                                break
                        except re.error as e:
                            print(f"Regex error for pattern '{regex_pattern}': {e}")
                            continue

            # If not found in paragraphs, try tables
            if not replaced:
                for table in doc.tables:
                    if replaced:
                        break

                    for row in table.rows:
                        if replaced:
                            break

                        for cell in row.cells:
                            if replaced:
                                break

                            for paragraph in cell.paragraphs:
                                if replaced:
                                    break

                                if paragraph.text:
                                    original_text = paragraph.text
                                    try:
                                        # Check if pattern exists in this paragraph
                                        if re.search(regex_pattern, original_text):
                                            # Replace only the FIRST occurrence
                                            new_text = re.sub(
                                                regex_pattern,
                                                replacement_value,
                                                original_text,
                                                count=1,
                                            )

                                            # Update paragraph and mark as replaced
                                            self._update_paragraph_text(
                                                paragraph, new_text
                                            )
                                            replacements_count += 1
                                            replaced = True
                                            break
                                    except re.error as e:
                                        print(
                                            f"Regex error for pattern '{regex_pattern}': {e}"
                                        )
                                        continue

        return replacements_count

    def _update_paragraph_text(self, paragraph, new_text: str):
        """
        Update paragraph text while preserving basic structure
        """
        # Clear existing runs and add new text
        # This preserves the paragraph but loses run-level formatting
        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


document_generator_service = DocumentGeneratorService()
